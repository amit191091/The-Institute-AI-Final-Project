"""
Document loaders for PDF, DOCX, and TXT files
"""
from pathlib import Path
from typing import List, Iterator, Tuple, Any
import logging

try:
    from unstructured.partition.pdf import partition_pdf  # type: ignore
    from unstructured.partition.docx import partition_docx  # type: ignore
    from unstructured.partition.text import partition_text  # type: ignore
except ImportError:
    logging.warning("Unstructured library not available. Some document types may not be supported.")
    partition_pdf = partition_docx = partition_text = None

try:
    import fitz  # type: ignore  # PyMuPDF for backup PDF processing
    PYMUPDF_AVAILABLE = True
except ImportError:
    fitz = None
    PYMUPDF_AVAILABLE = False
    logging.warning("PyMuPDF not available. PDF processing will use PyPDF2 fallback.")

try:
    import PyPDF2  # type: ignore
    PYPDF2_AVAILABLE = True
except ImportError:
    PyPDF2 = None
    PYPDF2_AVAILABLE = False

try:
    from docx import Document as DocxDocument  # type: ignore
    DOCX_AVAILABLE = True
except ImportError:
    DocxDocument = None
    DOCX_AVAILABLE = False

logger = logging.getLogger(__name__)

class DocumentLoader:
    """Enhanced document loader with fallback methods"""
    
    def __init__(self):
        self.use_unstructured = all([partition_pdf, partition_docx, partition_text])
        
    def load_elements(self, path: Path) -> List[Any]:
        """
        Load document elements with metadata preservation
        Returns list of elements from unstructured or custom parsers
        """
        if not path.exists():
            raise FileNotFoundError(f"Document not found: {path}")
            
        ext = path.suffix.lower()
        
        try:
            if self.use_unstructured:
                return self._load_with_unstructured(path, ext)
            else:
                return self._load_with_fallback(path, ext)
        except Exception as e:
            logger.error(f"Failed to load {path}: {e}")
            # Try fallback method
            return self._load_with_fallback(path, ext)
    
    def _load_with_unstructured(self, path: Path, ext: str) -> List[Any]:
        """Load using unstructured library (preferred)"""
        if ext == ".pdf":
            if partition_pdf is None:
                raise ImportError("unstructured.partition.pdf not available")
            return partition_pdf(filename=str(path), infer_table_structure=True)
        elif ext in (".docx", ".doc"):
            if partition_docx is None:
                raise ImportError("unstructured.partition.docx not available")
            return partition_docx(filename=str(path))
        elif ext == ".txt":
            if partition_text is None:
                raise ImportError("unstructured.partition.text not available")
            return partition_text(filename=str(path))
        else:
            raise ValueError(f"Unsupported format: {ext}")
    
    def _load_with_fallback(self, path: Path, ext: str) -> List[Any]:
        """Fallback document loading using basic libraries"""
        if ext == ".pdf":
            # Try PyMuPDF first, then PyPDF2
            if PYMUPDF_AVAILABLE:
                try:
                    return self._load_pdf_pymupdf(path)
                except Exception as e:
                    logger.warning(f"PyMuPDF failed: {e}, trying PyPDF2")
            
            if PYPDF2_AVAILABLE:
                return self._load_pdf_pypdf2(path)
            else:
                raise ImportError("No PDF processing libraries available")
                
        elif ext in (".docx", ".doc"):
            return self._load_docx_basic(path)
        elif ext == ".txt":
            return self._load_txt_basic(path)
        else:
            raise ValueError(f"Unsupported format: {ext}")
    
    def _load_pdf_pymupdf(self, path: Path) -> List[dict]:
        """Load PDF using PyMuPDF"""
        if not PYMUPDF_AVAILABLE:
            raise ImportError("PyMuPDF not available")
        elements: List[dict] = []
        # fitz is available per guard above
        assert fitz is not None
        doc = fitz.open(path)

        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            text = page.get_text()

            if text.strip():
                elements.append({
                    'text': text,
                    'category': 'Text',
                    'metadata': {'page_number': page_num + 1, 'id': f"page_{page_num + 1}"}
                })

            # Extract tables (basic)
            try:
                tables = page.find_tables()
            except Exception:
                tables = []
            for i, table in enumerate(tables):
                try:
                    table_text = table.extract()
                except Exception:
                    table_text = None
                if table_text:
                    table_str = "\\n".join([" | ".join(row) for row in table_text])
                    elements.append({
                        'text': table_str,
                        'category': 'Table',
                        'metadata': {'page_number': page_num + 1, 'id': f"table_{page_num + 1}_{i}"}
                    })

        doc.close()
        return elements

    def _load_pdf_pypdf2(self, path: Path) -> List[dict]:
        """Load PDF using PyPDF2 as fallback"""
        if not PYPDF2_AVAILABLE:
            raise ImportError("PyPDF2 not available")
            
        elements = []
        with open(path, 'rb') as file:
            assert PyPDF2 is not None
            reader = PyPDF2.PdfReader(file)
            
            for page_num, page in enumerate(reader.pages):
                text = page.extract_text()
                
                if text.strip():
                    elements.append({
                        'text': text,
                        'category': 'Text',
                        'metadata': {'page_number': page_num + 1, 'id': f"page_{page_num + 1}"}
                    })
        
        return elements
    
    def _load_docx_basic(self, path: Path) -> List[dict]:
        """Load DOCX using python-docx"""
        elements = []
        if DocxDocument is None:
            raise ImportError("python-docx not available")
        doc = DocxDocument(path)
        
        for i, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip():
                elements.append({
                    'text': paragraph.text,
                    'category': 'Text',
                    'metadata': {'page_number': 1, 'id': f"para_{i}"}
                })
        
        # Extract tables
        for i, table in enumerate(doc.tables):
            table_text = []
            for row in table.rows:
                row_text = [cell.text for cell in row.cells]
                table_text.append(" | ".join(row_text))
            
            if table_text:
                elements.append({
                    'text': "\\n".join(table_text),
                    'category': 'Table',
                    'metadata': {'page_number': 1, 'id': f"table_{i}"}
                })
        
        return elements
    
    def _load_txt_basic(self, path: Path) -> List[dict]:
        """Load TXT file with multiple encoding attempts"""
        encodings = ['utf-8', 'utf-8-sig', 'latin1', 'cp1252', 'iso-8859-1']
        
        for encoding in encodings:
            try:
                with open(path, 'r', encoding=encoding) as f:
                    text = f.read()
                
                return [{
                    'text': text,
                    'category': 'Text',
                    'metadata': {'page_number': 1, 'id': 'text_content', 'encoding': encoding}
                }]
            except UnicodeDecodeError:
                continue
        
        # If all encodings fail, read as binary and decode with errors='ignore'
        with open(path, 'rb') as f:
            raw_content = f.read()
            text = raw_content.decode('utf-8', errors='ignore')
        
        return [{
            'text': text,
            'category': 'Text',
            'metadata': {'page_number': 1, 'id': 'text_content', 'encoding': 'utf-8-ignore'}
        }]
    
    def load_many(self, paths: List[Path]) -> Iterator[Tuple[Path, List[Any]]]:
        """Load multiple documents"""
        for path in paths:
            try:
                elements = self.load_elements(path)
                yield path, elements
            except Exception as e:
                logger.error(f"Failed to load {path}: {e}")
                continue

def get_document_loader() -> DocumentLoader:
    """Get configured document loader instance"""
    return DocumentLoader()
